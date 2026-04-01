"""
Sudarshan Project Report — Word Document Generator
Part 1: Setup, helpers, and preliminary pages (cover through abstract)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAM_DIR = os.path.join(BASE_DIR, 'data', 'report_diagrams')
OUTPUT_PATH = os.path.join(BASE_DIR, 'Sudarshan_Project_Report.docx')

# ── Personal Details ──
STUDENT_NAME = 'Vedansh Gupta'
REG_NO = '<Registration No>'
PROJECT_TITLE = 'Sudarshan: An AI-Powered Web Vulnerability Scanner'
GUIDE_NAME = 'Ms. Anushka Bharadwaj'
HOD_NAME = 'Dr. Savita Shiwani'
TUTOR_NAME = 'Mrs Nitu Tank'
DEAN_NAME = 'Dr. Ajay Khunteta'
PRO_PRESIDENT = 'Dr. Manoj Gupta'

doc = Document()

# ═══════════════════════════════════════
# GLOBAL FORMATTING
# ═══════════════════════════════════════
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Set default font for whole doc via XML
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman" w:cs="Times New Roman"/>')
rPr.append(rFonts)


def set_margins(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.86)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(1.52)


def add_centered_text(text, size=12, bold=False, italic=False, color=None, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_justified_para(text, indent_first=True, bold=False, italic=False, size=12, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_chapter_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    # Add thick line after
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_after = Pt(12)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)
    return p


def add_heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    return p


def add_heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    return p


def add_heading3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p


def add_figure(image_path, caption_text, fig_num):
    """Add a centered figure with caption below."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.5))
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(f'Figure {fig_num}: {caption_text}')
        r.font.name = 'Garamond'
        r.font.size = Pt(10)
        r.bold = True


def add_table_with_data(headers, rows, table_num, caption):
    """Add a table with caption above."""
    # Caption above
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    r = cap.add_run(f'Table {table_num}: {caption}')
    r.font.name = 'Garamond'
    r.font.size = Pt(10)
    r.bold = True

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════
# SET MARGINS FOR DEFAULT SECTION
# ═══════════════════════════════════════
set_margins(doc.sections[0])

# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════
add_centered_text('FACULTY OF COMPUTER SCIENCE & ENGINEERING', 12, bold=True, space_after=2)
add_centered_text('DEPARTMENT OF COMPUTER APPLICATION', 12, bold=True, space_after=2)
add_centered_text('BACHELOR OF COMPUTER APPLICATION', 12, bold=True, space_after=2)
add_centered_text('(ACADEMIC SESSION: 2025-2026)', 12, bold=True, space_after=36)
add_centered_text('MAJOR PROJECT REPORT', 16, bold=True, space_after=12)
add_centered_text('ON', 14, bold=True, space_after=12)
add_centered_text(PROJECT_TITLE.upper(), 16, bold=True, space_after=48)
add_centered_text(STUDENT_NAME, 14, bold=True, space_after=4)
add_centered_text(f'{REG_NO} (3rd Year)', 12, space_after=48)

# Guide and HOD on same line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run(f'{GUIDE_NAME}')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p.add_run(f'                                        {HOD_NAME}')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(36)
r1 = p2.add_run('(Project Guide)')
r1.font.name = 'Times New Roman'
r1.font.size = Pt(11)
r2 = p2.add_run('                                              (Head of Department)')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(11)

add_centered_text('Poornima University', 13, bold=True, space_after=2)
add_centered_text('Plot No. 2027-2031, Ramchandrapura P.O. Vidhani Vatika Road,', 10, space_after=1)
add_centered_text('Sitapura, Jaipur, Rajasthan- 303905 (India)', 10, space_after=1)
add_centered_text('www.poornima.edu.in', 10, italic=True)

page_break()

# ═══════════════════════════════════════
# TITLE PAGE (FIRST PAGE)
# ═══════════════════════════════════════
add_centered_text('A', 14, bold=True, space_after=12, space_before=36)
add_centered_text('MAJOR PROJECT REPORT', 16, bold=True, space_after=6)
add_centered_text('On', 14, space_after=6)
add_centered_text(PROJECT_TITLE, 16, bold=True, space_after=24)
add_centered_text('Submitted in partial fulfilment of the requirements for the award of the Degree of', 12, space_after=12)
add_centered_text('BACHELOR OF COMPUTER APPLICATION', 14, bold=True, space_after=24)
add_centered_text('Poornima University, Jaipur (Academic Session: 2025-26)', 12, space_after=18)
add_centered_text('Submitted By:', 12, bold=True, space_after=4)
add_centered_text(STUDENT_NAME, 13, bold=True, space_after=2)
add_centered_text(REG_NO, 12, space_after=2)
add_centered_text('3rd Year, BCA', 12, space_after=24)
add_centered_text('Submitted To:', 12, bold=True, space_after=4)
add_centered_text('Department of Computer Application', 12, space_after=2)
add_centered_text('Faculty of Computer Science & Engineering, Poornima University', 12, space_after=2)
add_centered_text('Ramchandrapura, Sitapura Ext., Jaipur, Rajasthan- (303905)', 11)

page_break()

# ═══════════════════════════════════════
# CANDIDATE'S DECLARATION
# ═══════════════════════════════════════
add_chapter_title("CANDIDATE'S DECLARATION")
doc.add_paragraph()

declaration = (
    f'I hereby declare that the work presented in the Major Project report entitled '
    f'"{PROJECT_TITLE}" is submitted by {STUDENT_NAME} [{REG_NO}] in the fulfillment '
    f'of the requirements for the award of Bachelor of Computer Application in Faculty of '
    f'Computer Science & Engineering from Poornima University, Jaipur during the academic year '
    f'2025-26. The work has been found satisfactory, authentic of my own work carried out '
    f'during my degree and approved for submission.'
)
add_justified_para(declaration, indent_first=False)
add_justified_para(
    'The work reported in this has not been submitted by me for award of any other degree or diploma.',
    indent_first=False, space_before=6
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Date:')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run(STUDENT_NAME)
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.bold = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
r3 = p3.add_run('__.04.2026')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r4 = p4.add_run(f'[{REG_NO}]')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════
add_chapter_title('CERTIFICATE')
doc.add_paragraph()

cert_text = (
    f'This is to certify that the Major Project work entitled "{PROJECT_TITLE}" '
    f'is a bonafide work carried out in the VIth semester by {STUDENT_NAME} [{REG_NO}] '
    f'in partial fulfillment of the requirements for the degree of Bachelor of Computer '
    f'Application of Poornima University, Jaipur during the academic session 2025-26. '
    f'The project work has been found satisfactory and is approved for submission.'
)
add_justified_para(cert_text, indent_first=False)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run(GUIDE_NAME)
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r1.bold = True
r2 = p.add_run(f'                                        {HOD_NAME}')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p2.add_run('(Project Guide)')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(11)
r4 = p2.add_run('                                              (Head of Department)')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(11)

page_break()

# ═══════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════
add_chapter_title('ACKNOWLEDGEMENT')
doc.add_paragraph()

ack_paras = [
    (
        f'Undertaking this Major Project has been a meticulously planned and guided journey, '
        f'turning into a lifetime experience for me. Such an achievement would not have been possible '
        f'without the invaluable support from various sources and individuals at Poornima University.'
    ),
    (
        f'I extend my heartfelt gratitude to {DEAN_NAME}, Dean of the Faculty of Computer Science & '
        f'Engineering, for providing us with the platform and encouragement to successfully execute this project.'
    ),
    (
        f'I am deeply thankful to {HOD_NAME}, HoD of Computer Application, for her continuous support '
        f'and guidance throughout this endeavor. Her mentorship has been pivotal in navigating challenges '
        f'and completing this project successfully.'
    ),
    (
        f'Additionally, I extend my sincere appreciation to {PRO_PRESIDENT}, Pro-President of FCE, for his '
        f'encouragement and visionary guidance. His leadership and valuable inputs significantly contributed '
        f'to the overall success of our project.'
    ),
    (
        f'A special note of appreciation to my Project Guide {GUIDE_NAME} for her constant motivation, '
        f'technical expertise, and valuable insights throughout the project journey. Her guidance was '
        f'instrumental in shaping the direction and quality of this work.'
    ),
    (
        f'I am also grateful to our Class Tutor {TUTOR_NAME} for her consistent encouragement and support. '
        f'I extend my thanks to the librarian for providing access to essential resources and research materials '
        f'that were instrumental in our study.'
    ),
    (
        f'Furthermore, I extend my sincere gratitude to all the faculty members of the Department of Computer '
        f'Application for their unwavering support and guidance.'
    ),
    (
        f'Lastly, I deeply appreciate my friends and family, whose direct and indirect contributions, valuable '
        f'suggestions, and encouragement have played a vital role in the successful completion of this project.'
    ),
]

for i, para in enumerate(ack_paras):
    add_justified_para(para, indent_first=(i > 0), space_after=4)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(STUDENT_NAME)
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run(f'[{REG_NO}]')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════
add_chapter_title('ABSTRACT')
doc.add_paragraph()
doc.add_paragraph()

abstract_paras = [
    (
        'The escalating sophistication of cyber threats demands equally advanced defensive tooling. '
        'This project presents Sudarshan, a full-stack web vulnerability scanner engineered to bridge the '
        'gap between conventional automated scanners and the nuanced reasoning of human penetration testers. '
        'Built on Python 3.12 and the Flask 3.0 framework, Sudarshan integrates a multi-threaded crawling engine '
        'with sixteen specialized vulnerability detection modules spanning the OWASP Top 10 and beyond, covering '
        'SQL Injection, XSS, CSRF, SSRF, SSTI, XXE, JWT attacks, and several other critical vulnerability classes.'
    ),
    (
        'What distinguishes Sudarshan from existing scanners is its three-layered AI intelligence system. '
        'At its core, the SmartEngine unifies a large language model (Groq-hosted Llama 3.3 70B) with a curated '
        'PortSwigger knowledge base of 269 labs and 2,197 payloads, and a machine-learning false-positive classifier '
        'built on a Random Forest and Gradient Boosting ensemble. The LLM performs target reconnaissance, generates '
        'context-aware payloads, crafts WAF bypass strategies, and produces professional attack narratives. '
        'Simultaneously, the ML classifier filters false positives using sixteen engineered features, and its verdict '
        'is blended with the LLM assessment in a weighted 40/60 scheme to deliver high-confidence results.'
    ),
    (
        'The platform supports multi-tenant organizations with role-based access, HMAC-SHA256 API key authentication, '
        'real-time Server-Sent Event streaming, and containerized deployment via Docker with Celery-based background task '
        'processing. Scan reports include AI-generated executive summaries, remediation plans with code examples, '
        'and detailed attack narratives enriched with PortSwigger Academy references. The system has been tested against '
        'DVWA and other intentionally vulnerable applications, successfully identifying critical vulnerabilities '
        'while demonstrating measurable reductions in false-positive rates through the combined AI verification pipeline.'
    ),
]

for para in abstract_paras:
    add_justified_para(para, indent_first=False, italic=True, space_after=6)

page_break()

# Save what we have so far — the main chapters will be added next
# We'll save the doc object for the next script to pick up
import pickle
with open(os.path.join(BASE_DIR, 'scripts', '_report_state.pkl'), 'wb') as f:
    pickle.dump({
        'doc_path': OUTPUT_PATH,
        'diagram_dir': DIAGRAM_DIR,
        'project_title': PROJECT_TITLE,
        'student_name': STUDENT_NAME,
        'reg_no': REG_NO,
    }, f)

doc.save(OUTPUT_PATH)
print(f'Part 1 saved: {OUTPUT_PATH}')
