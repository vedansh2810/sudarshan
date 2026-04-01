"""
Sudarshan Project Report — Word Document Generator
Part 2: Chapters 1-3 (Introduction, Literature Survey, System Design)
Run AFTER generate_report_p1.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAM_DIR = os.path.join(BASE_DIR, 'data', 'report_diagrams')
OUTPUT_PATH = os.path.join(BASE_DIR, 'Sudarshan_Project_Report.docx')

PROJECT_TITLE = 'Sudarshan: An AI-Powered Web Vulnerability Scanner'

doc = Document(OUTPUT_PATH)

# ── Helpers (duplicated for standalone execution) ──
def add_centered_text(text, size=12, bold=False, italic=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_justified_para(text, indent_first=True, bold=False, italic=False, size=12, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_chapter_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_after = Pt(12)
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

def add_figure(image_path, caption_text, fig_num):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(f'Figure {fig_num}: {caption_text}')
        r.font.name = 'Garamond'
        r.font.size = Pt(10)
        r.bold = True

def add_table_with_data(headers, rows, table_num, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    r = cap.add_run(f'Table {table_num}: {caption}')
    r.font.name = 'Garamond'
    r.font.size = Pt(10)
    r.bold = True
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    bullet = '\u2022' if level == 0 else '\u25E6'
    run = p.add_run(f'{bullet}  {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════
add_chapter_title('1. INTRODUCTION')

add_heading1('1.1 PROJECT OVERVIEW')
add_justified_para(
    'The rapid proliferation of web-based applications across industries — from banking and healthcare to '
    'e-commerce and government services — has made web security one of the most critical concerns in modern '
    'software engineering. Every web application deployed to the internet becomes a potential target for attackers '
    'who exploit vulnerabilities in code, configuration, or architecture to steal data, disrupt services, or '
    'gain unauthorized access. The OWASP Foundation reports that injection flaws, broken authentication, and '
    'cross-site scripting remain persistently among the top risks, despite decades of awareness campaigns and '
    'tooling development.',
    indent_first=False
)
add_justified_para(
    'Sudarshan is a comprehensive, AI-powered web vulnerability scanner designed and built from the ground up '
    'to address the limitations of conventional scanning tools. Named after the mythological divine weapon '
    'symbolizing precision and protection, Sudarshan employs a multi-layered approach that combines traditional '
    'vulnerability scanning techniques with modern artificial intelligence to deliver faster, smarter, and more '
    'accurate security assessments.'
)
add_justified_para(
    'Unlike commercial scanners that rely solely on signature matching and predefined payloads, Sudarshan integrates '
    'a large language model for intelligent reasoning about vulnerabilities, a curated knowledge base from '
    'PortSwigger Academy for expert-level payloads, and a machine learning classifier to minimize false positives. '
    'The result is a tool that thinks more like a human penetration tester while maintaining the speed and '
    'thoroughness of automated scanning.'
)

add_heading1('1.2 MOTIVATION')
add_justified_para(
    'During our coursework and participation in security-focused labs, we noticed a significant gap between '
    'the capabilities of freely available scanning tools and the depth of analysis performed by experienced '
    'penetration testers. Tools such as OWASP ZAP and Nikto are excellent starting points, but they often '
    'produce high volumes of false positives, lack contextual awareness of the target application, and provide '
    'generic remediation advice that developers find difficult to act upon.',
    indent_first=False
)
add_justified_para(
    'The emergence of powerful large language models in 2023-2024 presented a unique opportunity: what if the '
    'reasoning capabilities of an LLM could be combined with structured security knowledge and statistical '
    'machine learning to create a scanner that not only finds vulnerabilities but understands them? This question '
    'became the driving motivation behind Sudarshan.'
)
add_justified_para(
    'Furthermore, the increasing adoption of AI by attackers — for crafting sophisticated phishing campaigns, '
    'generating polymorphic malware, and automating exploitation — necessitates that defensive tools evolve at '
    'the same pace. Sudarshan represents an effort to bring AI capabilities to the defenders\' side of the equation.'
)

add_heading1('1.3 OBJECTIVES')
add_justified_para('The primary objectives of this project are as follows:', indent_first=False)
add_bullet('To design and implement a full-stack web vulnerability scanner capable of detecting sixteen distinct vulnerability types conforming to OWASP guidelines.')
add_bullet('To integrate a large language model (Llama 3.3 70B) for context-aware reconnaissance, payload generation, and vulnerability analysis.')
add_bullet('To build a machine-learning false-positive classifier using Random Forest and Gradient Boosting ensemble techniques.')
add_bullet('To incorporate a curated PortSwigger knowledge base with 269 labs and 2,197 payloads for expert-level attack strategies.')
add_bullet('To develop a multi-tenant SaaS architecture supporting organizations, role-based access, and API key authentication.')
add_bullet('To generate AI-powered security reports with executive summaries, remediation plans, and attack narratives.')
add_bullet('To containerize the entire application stack using Docker for reproducible deployment.')

add_heading1('1.4 SCOPE OF THE PROJECT')
add_justified_para(
    'Sudarshan is scoped as a black-box web application vulnerability scanner. It operates by sending HTTP '
    'requests to the target application, analyzing responses, and inferring the presence of vulnerabilities '
    'without requiring access to the application source code. The scanner supports both active scanning '
    '(injecting payloads) and passive scanning (analyzing headers and configurations).',
    indent_first=False
)
add_justified_para(
    'The project encompasses the complete software development lifecycle: requirements analysis, system design, '
    'implementation, testing, and deployment. It includes a web-based user interface for scan management, '
    'a RESTful API for programmatic access, and Dockerized deployment for production environments. The scope '
    'explicitly excludes network-layer scanning, mobile application testing, and source-code analysis, which '
    'belong to different categories of security tools.'
)

page_break()

# ═══════════════════════════════════════════
# CHAPTER 2: LITERATURE SURVEY
# ═══════════════════════════════════════════
add_chapter_title('2. LITERATURE SURVEY')

add_heading1('2.1 WEB APPLICATION SECURITY LANDSCAPE')
add_justified_para(
    'Web application security has matured significantly since the early days of the internet. The OWASP '
    'Foundation, established in 2001, has been instrumental in cataloguing and categorizing web vulnerabilities. '
    'Their periodic OWASP Top 10 list serves as the de facto standard for web security awareness, with the '
    '2021 edition highlighting Broken Access Control, Cryptographic Failures, and Injection as the three most '
    'critical risk categories.',
    indent_first=False
)
add_justified_para(
    'The Common Weakness Enumeration (CWE) maintained by MITRE provides a more granular taxonomy with over '
    '900 weakness types, while the Common Vulnerability Scoring System (CVSS) offers a standardized method '
    'for rating the severity of vulnerabilities on a 0-10 scale. These frameworks form the backbone of '
    'Sudarshan\'s vulnerability classification and scoring mechanisms.'
)

add_heading1('2.2 EXISTING VULNERABILITY SCANNING TOOLS')
add_justified_para(
    'Several tools exist in the web vulnerability scanning space, each with distinct strengths and limitations. '
    'Understanding these tools was essential to identifying the gaps Sudarshan aims to fill.',
    indent_first=False
)

add_heading2('2.2.1 Open-Source Scanners')
add_justified_para(
    'OWASP ZAP (Zed Attack Proxy) is arguably the most widely used open-source scanner. It functions as an '
    'intercepting proxy with active and passive scanning capabilities. While feature-rich and extensible through '
    'add-ons, ZAP relies primarily on pattern matching and lacks contextual intelligence about the target '
    'application. Its false-positive rates on complex applications can exceed 30%, requiring significant manual '
    'review effort.',
    indent_first=False
)
add_justified_para(
    'Nikto is a lightweight web server scanner focused on misconfigurations and known vulnerabilities in server '
    'software. While fast and reliable for its intended purpose, it does not perform application-level testing '
    'such as SQL injection or XSS detection. SQLMap specializes exclusively in SQL injection detection and '
    'exploitation, offering deep coverage for that single vulnerability class but providing no capability for '
    'other attack vectors.'
)

add_heading2('2.2.2 Commercial Scanners')
add_justified_para(
    'Commercial products such as Burp Suite Professional, Acunetix, and Invicti (formerly Netsparker) offer '
    'more sophisticated scanning engines with lower false-positive rates. Burp Suite, in particular, is the '
    'industry standard for manual penetration testing, and its scanner component benefits from years of research '
    'by PortSwigger. However, these tools carry significant licensing costs — Burp Suite Professional starts at '
    '$449/year per user — making them inaccessible for students, small organizations, and developing nations.',
    indent_first=False
)

add_heading1('2.3 ARTIFICIAL INTELLIGENCE IN CYBERSECURITY')
add_justified_para(
    'The application of artificial intelligence to cybersecurity has followed two parallel tracks: machine '
    'learning for pattern recognition and anomaly detection, and more recently, large language models for '
    'reasoning about security contexts.',
    indent_first=False
)

add_heading2('2.3.1 Machine Learning Approaches')
add_justified_para(
    'Traditional ML approaches in vulnerability detection have focused on binary classification problems: '
    'determining whether a given scanner finding is a true positive or false positive. Research by Arp et al. '
    '(2022) demonstrated that ensemble methods — particularly Random Forests and Gradient Boosting — achieve '
    'superior performance on security classification tasks due to their ability to handle imbalanced datasets '
    'and heterogeneous feature types. Sudarshan\'s false-positive classifier directly implements this finding, '
    'using a Random Forest and Gradient Boosting ensemble with sixteen engineered features covering payload '
    'characteristics, HTTP response analysis, and error pattern detection.',
    indent_first=False
)

add_heading2('2.3.2 Large Language Models in Security')
add_justified_para(
    'The release of GPT-3 in 2020 and subsequent models demonstrated that LLMs could understand and generate '
    'technical content with remarkable proficiency. In the security domain, researchers have explored LLMs for '
    'vulnerability explanation, code review, and even exploit generation. Pearce et al. (2023) showed that LLMs '
    'can identify security vulnerabilities in code with accuracy comparable to static analysis tools, while '
    'providing contextual explanations that developers find more actionable.',
    indent_first=False
)
add_justified_para(
    'Sudarshan leverages the Llama 3.3 70B model via the Groq inference API for multiple security tasks: '
    'target reconnaissance, payload generation, WAF bypass strategy development, vulnerability analysis, and '
    'professional report writing. The key architectural decision was to make AI features supplementary — the '
    'scanner operates fully even when the LLM is unavailable, ensuring reliability while benefiting from AI '
    'intelligence when available.'
)

add_heading1('2.4 PORTSWIGGER WEB SECURITY ACADEMY')
add_justified_para(
    'PortSwigger\'s Web Security Academy represents the most comprehensive publicly available collection of '
    'web security knowledge. With over 269 interactive labs covering 31 vulnerability categories, the Academy '
    'provides not just theoretical knowledge but practical exploitation techniques with working payloads. '
    'Sudarshan incorporates a structured knowledge base scraped and indexed from this resource, enabling the '
    'SmartEngine to reference real-world attack techniques and map findings to specific labs for educational value.',
    indent_first=False
)

add_heading1('2.5 IDENTIFIED GAPS AND CONTRIBUTION')
add_justified_para('Based on the above literature survey, the following gaps were identified:', indent_first=False)
add_bullet('No open-source scanner combines LLM reasoning with structured knowledge bases and ML classification in a unified architecture.')
add_bullet('Existing tools treat AI as an afterthought rather than a core component of the scanning pipeline.')
add_bullet('False-positive filtering in most scanners is rule-based rather than learned from data.')
add_bullet('Report generation is typically template-based, producing generic outputs that lack contextual depth.')
add_justified_para(
    'Sudarshan directly addresses each of these gaps through its SmartEngine architecture, which tightly '
    'integrates LLM, knowledge base, and ML classifier into every phase of the scanning pipeline.',
    space_before=6
)

page_break()

# ═══════════════════════════════════════════
# CHAPTER 3: SYSTEM DESIGN AND ARCHITECTURE
# ═══════════════════════════════════════════
add_chapter_title('3. SYSTEM DESIGN AND ARCHITECTURE')

add_heading1('3.1 TECHNOLOGY STACK')
add_justified_para(
    'Sudarshan\'s technology choices were driven by three principles: developer productivity, ecosystem maturity, '
    'and AI integration capabilities. The full stack is summarized below.',
    indent_first=False
)

add_table_with_data(
    ['Layer', 'Technology', 'Version', 'Purpose'],
    [
        ['Backend Framework', 'Flask', '3.0.0', 'Lightweight, extensible web framework'],
        ['Language', 'Python', '3.12', 'Rich security and AI library ecosystem'],
        ['Database', 'PostgreSQL (Supabase)', '15+', 'Production-grade relational storage'],
        ['ORM', 'Flask-SQLAlchemy', '3.1.1', 'Declarative database modeling'],
        ['Authentication', 'Supabase Auth (GoTrue)', '2.0+', 'JWT-based user authentication'],
        ['AI / LLM', 'Groq API (Llama 3.3 70B)', '-', 'Intelligent reasoning and analysis'],
        ['Machine Learning', 'scikit-learn', '1.3.2', 'False-positive classification'],
        ['Task Queue', 'Celery + Redis', '5.4.0', 'Async background scan processing'],
        ['Containerization', 'Docker + docker-compose', '-', 'Reproducible multi-service deployment'],
        ['Report Generation', 'fpdf2', '2.7.6', 'PDF report export'],
    ],
    '3.1', 'Technology Stack Overview'
)

add_heading1('3.2 SYSTEM ARCHITECTURE')
add_justified_para(
    'The system follows a layered architecture pattern with clear separation of concerns between the '
    'presentation layer, application logic, scanner engine, AI intelligence layer, and data persistence. '
    'This design enables independent evolution of each layer and facilitates horizontal scaling through '
    'containerization.',
    indent_first=False
)
add_justified_para(
    'At the topmost level, users interact with the system through either the web interface (server-rendered '
    'Jinja2 templates) or the RESTful API v2. Authentication flows through Supabase Auth, which handles '
    'registration, login, and JWT token management. The Flask application routes requests to the appropriate '
    'blueprints, which in turn delegate scanning operations to the Scanner Engine.'
)
add_justified_para(
    'The Scanner Engine operates as a semi-autonomous subsystem. Upon receiving a scan request, the Scan '
    'Manager decides whether to execute the scan in a Celery background worker (if Redis is available) or '
    'in an in-process thread (for development environments). This dual-mode execution ensures the system '
    'functions correctly regardless of infrastructure availability.'
)

add_figure(os.path.join(DIAGRAM_DIR, 'system_architecture.png'),
           'Sudarshan System Architecture', '3.1')

add_heading1('3.3 DATABASE DESIGN')
add_justified_para(
    'The database schema was designed to support multi-tenant operation, comprehensive scan data storage, '
    'and ML training data collection. The schema comprises eleven interrelated tables organized around '
    'three core domains: user management, scan operations, and machine learning.',
    indent_first=False
)

add_heading2('3.3.1 Core Tables')
add_justified_para(
    'The Users table stores local user records mapped from Supabase Auth via the supabase_uid foreign key. '
    'This dual-identity approach allows the application to maintain its own user metadata while leveraging '
    'Supabase for secure authentication. Each user can belong to multiple Organizations through the '
    'OrgMemberships junction table, supporting roles of owner, admin, member, and viewer.',
    indent_first=False
)
add_justified_para(
    'The Scans table captures comprehensive metadata about each scanning operation, including target URL, '
    'scan mode (active/passive), scan speed configuration, progress metrics (total URLs discovered, URLs '
    'tested, vulnerability counts by severity), and the final security score. Scans are scoped to both a '
    'user and optionally an organization, enabling shared visibility within teams.'
)
add_justified_para(
    'The Vulnerabilities table stores each detected finding with rich contextual data: vulnerability type, '
    'severity (critical/high/medium/low/info), CVSS score, OWASP category mapping, the specific payload '
    'that triggered the finding, complete request and response data for evidence, AI-generated analysis, '
    'attack narratives, and false-positive confidence scores from both the ML classifier and LLM verifier.'
)

add_figure(os.path.join(DIAGRAM_DIR, 'er_diagram.png'),
           'Database Entity-Relationship Diagram', '3.2')

add_heading2('3.3.2 Supporting Tables')
add_justified_para(
    'CrawledUrls maintains a record of every URL discovered during the crawling phase, along with HTTP '
    'status codes and counts of forms and parameters found at each URL. ScanLogs provides a chronological '
    'event log for each scan, enabling real-time progress tracking through SSE. Webhooks allows users to '
    'register HTTP callback URLs that receive notifications on scan completion, vulnerability discovery, '
    'or scan errors. APIKeys supports programmatic access through HMAC-SHA256 hashed tokens with '
    'configurable expiration dates.',
    indent_first=False
)

add_heading2('3.3.3 ML Training Tables')
add_justified_para(
    'ScanAttempts records every individual payload test during scanning, capturing sixteen engineered '
    'features for each attempt. These features include payload length, encoding type, parameter position, '
    'response status code, response time, content length delta, and error pattern matches. When findings are '
    'manually labeled as true or false positives, this data serves as the training set for the ML '
    'false-positive classifier. The MLModels table tracks trained model versions with their accuracy metrics, '
    'enabling model versioning and A/B comparisons.',
    indent_first=False
)

add_heading1('3.4 SCAN PIPELINE DESIGN')
add_justified_para(
    'The scan pipeline follows a four-phase architecture designed for thoroughness, intelligence, and '
    'performance. Each phase builds upon the outputs of the previous phase, creating an increasingly rich '
    'understanding of the target application.',
    indent_first=False
)

add_figure(os.path.join(DIAGRAM_DIR, 'scan_pipeline.png'),
           'Scan Pipeline Data Flow', '3.3')

add_heading2('3.4.1 Phase 0 — Connectivity Check')
add_justified_para(
    'Before committing resources to a full scan, the system performs a lightweight HTTP GET request to the '
    'target URL to verify reachability. This phase checks for DNS resolution errors, connection timeouts, '
    'SSL certificate issues, and HTTP redirect chains. If the target is unreachable, the scan terminates '
    'with a descriptive error message rather than consuming resources unnecessarily.',
    indent_first=False
)

add_heading2('3.4.2 Phase 1 — Web Crawling')
add_justified_para(
    'The multi-threaded crawler systematically discovers the target application\'s attack surface. Starting '
    'from the seed URL, it follows links, parses HTML forms, identifies URL parameters, and maps the '
    'application\'s structure. The crawler respects configurable depth limits and URL caps based on the '
    'selected scan speed profile. Thread counts range from 3 (safe mode) to 10 (aggressive mode), with '
    'inter-request delays from 1.0 second to 0.05 seconds respectively.',
    indent_first=False
)

add_heading2('3.4.3 Phase 1.5 — AI Reconnaissance')
add_justified_para(
    'Unique to Sudarshan, this intermediate phase uses the LLM to analyze initial HTTP responses and '
    'identify the target\'s technology stack, web application firewall, and framework. This intelligence '
    'shapes subsequent scanning decisions: if a WAF is detected, the system prioritizes bypass payloads; '
    'if a specific framework is identified, framework-specific vulnerability checks are weighted more heavily. '
    'This mimics the reconnaissance methodology of human penetration testers.',
    indent_first=False
)

add_heading2('3.4.4 Phase 2 — Vulnerability Scanning')
add_justified_para(
    'The sixteen vulnerability scanner modules execute in parallel via Python\'s ThreadPoolExecutor. Each '
    'module implements a base scanner interface and specializes in a specific vulnerability class. For every '
    'finding, the pipeline triggers four sequential operations: database persistence, AI-powered analysis '
    '(OWASP/CWE mapping and explanation), false-positive verification (ML 40% + LLM 60% combined verdict), '
    'and attack narrative generation. Progress updates stream to the frontend in real-time via Server-Sent '
    'Events.',
    indent_first=False
)

add_heading2('3.4.5 Phase 3 — Post-Scan AI Analysis')
add_justified_para(
    'After all scanners complete, the SmartEngine performs a holistic analysis of critical and high-severity '
    'findings. This phase generates the executive summary, creates a prioritized remediation plan with code '
    'examples, produces detailed attack narratives enriched with PortSwigger Academy references, and '
    'calculates the overall security risk score with business-friendly explanations.',
    indent_first=False
)

add_heading1('3.5 AUTHENTICATION AND AUTHORIZATION DESIGN')
add_justified_para(
    'The authentication architecture employs a hybrid model combining Supabase Auth for identity management '
    'with local session-based access control. User registration and login are handled by Supabase\'s GoTrue '
    'service, which manages password hashing, email verification, and JWT token generation. Upon successful '
    'authentication, the application creates a local user record mapped via the supabase_uid foreign key.',
    indent_first=False
)
add_justified_para(
    'Web interface access is controlled by the @login_required decorator, which validates the session cookie '
    'and retrieves the current user\'s permissions. API access uses HMAC-SHA256 hashed API keys passed via '
    'the X-API-Key header. Rate limiting is enforced globally at 200 requests per day and 50 per hour, with '
    'configurable overrides per endpoint. CSRF protection is enabled for all form submissions through '
    'Flask-WTF with a one-hour token validity period.'
)

page_break()
doc.save(OUTPUT_PATH)
print(f'Part 2 saved (Chapters 1-3): {OUTPUT_PATH}')
